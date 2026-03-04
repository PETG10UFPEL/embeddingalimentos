# ingest.py
"""
Indexador (RAG) para o projeto PET-Saúde G10.
Usa Google Generative AI Embeddings (text-embedding-004).

Uso:
  python ingest.py
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import List, Tuple

from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma
from docx import Document as DocxDocument

# Carrega as variáveis do arquivo .env (Certifique-se que a GOOGLE_API_KEY está lá)
load_dotenv()

# Configurações de diretório e coleção
RAW_DIR_DEFAULT = "data/raw_docs"
DB_DIR_DEFAULT = "data/chroma_db"
COLLECTION_NAME = "diet_knowledge"

def _load_pdf(path: Path) -> List[Document]:
    """Carrega arquivos PDF."""
    return PyPDFLoader(str(path)).load()

def _load_docx(path: Path) -> List[Document]:
    """Carrega arquivos DOCX extraindo parágrafos e tabelas."""
    doc = DocxDocument(str(path))
    parts: List[str] = []
    
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
            
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
                
    text = "\n".join(parts).strip()
    if not text:
        return []
    return [Document(page_content=text, metadata={"source": str(path), "page": "DOCX"})]

def load_file(path: Path) -> List[Document]:
    """Identifica a extensão e chama o carregador apropriado."""
    suf = path.suffix.lower()
    if suf == ".pdf":
        return _load_pdf(path)
    if suf == ".docx":
        return _load_docx(path)
    return []

def load_all_docs(raw_dir: str) -> Tuple[List[Document], List[str]]:
    """Varre a pasta de documentos e carrega todos os arquivos válidos."""
    raw = Path(raw_dir)
    raw.mkdir(parents=True, exist_ok=True)
    docs: List[Document] = []
    skipped: List[str] = []
    
    for p in raw.rglob("*"):
        if not p.is_file():
            continue
        if p.suffix.lower() not in [".pdf", ".docx"]:
            skipped.append(str(p))
            continue
        try:
            docs.extend(load_file(p))
        except Exception as e:
            skipped.append(f"{p} (erro: {e})")
    return docs, skipped

def build_index(
    raw_dir: str = RAW_DIR_DEFAULT,
    db_dir: str = DB_DIR_DEFAULT,
    chunk_size: int = 900,
    chunk_overlap: int = 150,
) -> int:
    """Processo principal de indexação."""
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise RuntimeError("GOOGLE_API_KEY não encontrada no seu arquivo .env.")

    # 1. CARREGAMENTO (Correção do NameError: docs)
    docs, skipped = load_all_docs(raw_dir)

    if not docs:
        print(f"ERRO: Nenhum documento válido (.pdf ou .docx) encontrado em {raw_dir}")
        return 0

    # 2. DIVISÃO EM CHUNKS
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(docs)

    # 3. CONFIGURAÇÃO DO MODELO (Sincronizado com rag.py)
    # Removido o prefixo 'models/' para evitar erro 404 na API v1beta
    # Tente usar o ID estável sem o prefixo 'models/'
    # Tente o modelo mais estável da frota Google
    # Use o modelo que o seu diagnóstico confirmou como disponível
    embeddings = GoogleGenerativeAIEmbeddings(
    model="models/gemini-embedding-001", 
    google_api_key=os.getenv("GOOGLE_API_KEY")
    )
    # 4. CRIAÇÃO DO BANCO VETORIAL (CHROMA)
    # Lembrete: Apague a pasta data/chroma_db manualmente antes de rodar
    vectordb = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=db_dir,
        collection_name=COLLECTION_NAME,
    )

    if skipped:
        print(f"\nArquivos ignorados ou com erro: {len(skipped)}")

    print(f"\nSucesso, Alejandro! {len(chunks)} trechos indexados em: {db_dir}")
    return len(chunks)

if __name__ == "__main__":
    build_index()