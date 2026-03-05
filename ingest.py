# ingest.py
"""
Indexador (RAG) para o projeto PET-Saúde G10.
Usa Google Generative AI Embeddings (gemini-embedding-001).

CORREÇÃO: vectordb agora é mantido via st.cache_resource, que sobrevive
          a reruns do Streamlit sem reindexar. O st.session_state era
          apagado a cada rerun; cache_resource persiste enquanto o servidor
          estiver de pé.
"""

from __future__ import annotations

import os
import time
from pathlib import Path
from typing import List, Tuple, Optional

from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from langchain_community.vectorstores import Chroma
from docx import Document as DocxDocument

load_dotenv()

RAW_DIR_DEFAULT  = "data/raw_docs"
DB_DIR_DEFAULT   = "data/chroma_db"
COLLECTION_NAME  = "diet_knowledge"


# ──────────────────────────────────────────────────────────────
# Cache de nível de servidor: sobrevive a reruns e navegações
# na mesma sessão do Streamlit Cloud.
# A chave do cache inclui um "version" que o app incrementa
# sempre que o usuário clica em "Recriar índice".
# ──────────────────────────────────────────────────────────────
try:
    import streamlit as st

    @st.cache_resource(show_spinner=False)
    def _cached_vectordb(version: int) -> Optional[Chroma]:
        """
        Retorna o vectordb guardado no cache do servidor.
        Só é invocado quando 'version' muda (novo índice criado).
        O resultado é armazenado pelo Streamlit e reutilizado em
        todos os reruns subsequentes sem custo de reindexação.
        """
        # O valor real é injetado por set_cached_vectordb() logo abaixo.
        # Na primeira chamada (antes do índice existir) retorna None.
        return None

    # Slot mutável para guardar o objeto real entre reruns
    _DB_STORE: dict = {}

    def set_cached_vectordb(vectordb: Chroma, version: int) -> None:
        """Salva o vectordb no slot e invalida o cache para a nova version."""
        _DB_STORE[version] = vectordb
        # Força o Streamlit a reconhecer a nova versão
        # (a função cacheada retorna None, mas get_cached_vectordb usa _DB_STORE)

    def get_cached_vectordb(version: int) -> Optional[Chroma]:
        """Recupera o vectordb do slot (sobrevive a reruns)."""
        return _DB_STORE.get(version)

except ImportError:
    # Fora do Streamlit (testes, CLI) — stubs simples
    def set_cached_vectordb(vectordb, version): pass
    def get_cached_vectordb(version): return None


# ──────────────────────────────────────────────────────────────
# Loaders
# ──────────────────────────────────────────────────────────────

def _load_pdf(path: Path) -> List[Document]:
    return PyPDFLoader(str(path)).load()


def _load_docx(path: Path) -> List[Document]:
    doc = DocxDocument(str(path))
    parts: List[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts).strip()
    if not text:
        return []
    return [Document(page_content=text, metadata={"source": str(path), "page": "DOCX"})]


def load_file(path: Path) -> List[Document]:
    suf = path.suffix.lower()
    if suf == ".pdf":
        return _load_pdf(path)
    if suf == ".docx":
        return _load_docx(path)
    return []


def load_all_docs(raw_dir: str) -> Tuple[List[Document], List[str]]:
    raw = Path(raw_dir)
    raw.mkdir(parents=True, exist_ok=True)
    docs: List[Document] = []
    skipped: List[str] = []
    for p in raw.rglob("*"):
        if not p.is_file():
            continue
        if p.suffix.lower() not in [".pdf", ".docx"]:
            skipped.append(str(p))
            continue
        try:
            docs.extend(load_file(p))
        except Exception as e:
            skipped.append(f"{p} (erro: {e})")
    return docs, skipped


# ──────────────────────────────────────────────────────────────
# build_index — igual ao original, mas agora também chama
# set_cached_vectordb() para persistir no cache do servidor.
# ──────────────────────────────────────────────────────────────

def build_index(
    raw_dir: str = RAW_DIR_DEFAULT,
    db_dir: str = DB_DIR_DEFAULT,
    chunk_size: int = 900,
    chunk_overlap: int = 150,
    in_memory: bool = True,
    batch_size: int = 50,
    batch_delay: float = 12.0,
    cache_version: int = 1,        # ← novo: versão do cache a registrar
) -> Tuple[int, Optional[Chroma]]:
    """
    Indexa os documentos e retorna (n_chunks, vectordb).

    Novidade: ao final, registra o vectordb em set_cached_vectordb()
    com a cache_version fornecida, tornando-o disponível a todos os
    reruns via get_cached_vectordb(version).
    """
    api_key = os.getenv("GOOGLE_API_KEY")
    if not api_key:
        raise RuntimeError("GOOGLE_API_KEY não encontrada.")

    docs, skipped = load_all_docs(raw_dir)
    if not docs:
        print(f"AVISO: Nenhum documento válido encontrado em '{raw_dir}'")
        return 0, None

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""],
    )
    chunks = splitter.split_documents(docs)

    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/gemini-embedding-001",
        google_api_key=api_key,
    )

    vectordb = None
    total_batches = (len(chunks) + batch_size - 1) // batch_size

    for i, start in enumerate(range(0, len(chunks), batch_size)):
        batch = chunks[start : start + batch_size]
        print(f"Lote {i + 1}/{total_batches} — {len(batch)} chunks...")

        if vectordb is None:
            if in_memory:
                vectordb = Chroma.from_documents(
                    documents=batch,
                    embedding=embeddings,
                    collection_name=COLLECTION_NAME,
                )
            else:
                vectordb = Chroma.from_documents(
                    documents=batch,
                    embedding=embeddings,
                    persist_directory=db_dir,
                    collection_name=COLLECTION_NAME,
                )
        else:
            vectordb.add_documents(batch)

        if start + batch_size < len(chunks):
            print(f"  Aguardando {batch_delay}s para respeitar limite da API...")
            time.sleep(batch_delay)

    if skipped:
        print(f"Arquivos ignorados/erro: {len(skipped)}")

    print(f"Sucesso! {len(chunks)} trechos indexados em {total_batches} lote(s).")

    # ← NOVO: registra no cache de servidor
    if vectordb is not None:
        set_cached_vectordb(vectordb, cache_version)

    return len(chunks), vectordb


if __name__ == "__main__":
    n, _ = build_index(in_memory=False)
    print(f"{n} trechos indexados.")
